from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path
from typing import Any, List, Tuple

import pandas as pd
from openpyxl import load_workbook
from docx import Document
from docx.table import Table


# ----------------------------
# Utilities
# ----------------------------
def is_empty(v: Any) -> bool:
    if v is None:
        return True
    if isinstance(v, str) and v.strip() == "":
        return True
    return False


def normalize_cell(v: Any) -> str:
    if v is None:
        return ""
    # openpyxl は日付を datetime/date にすることが多い
    if hasattr(v, "isoformat"):
        try:
            return v.isoformat(sep=" ")
        except TypeError:
            return str(v)
    return str(v)


def sheet_used_bounds(values: List[List[Any]]) -> Tuple[int, int, int, int]:
    """
    values: 2D list [row][col] (0-based)
    Returns: (r0, r1, c0, c1) inclusive bounds of non-empty cells.
    If all empty -> (0, -1, 0, -1)
    """
    non_empty = []
    for r, row in enumerate(values):
        for c, v in enumerate(row):
            if is_empty(v):
                continue
            non_empty.append((r, c))
    if not non_empty:
        return (0, -1, 0, -1)
    rs = [rc[0] for rc in non_empty]
    cs = [rc[1] for rc in non_empty]
    return (min(rs), max(rs), min(cs), max(cs))


def safe_filename(name: str, max_len: int = 80) -> str:
    # Windows禁止文字を置換
    name = re.sub(r'[\\/:*?"<>|]', "_", name)
    name = re.sub(r"\s+", " ", name).strip()
    if not name:
        name = "sheet"
    return name[:max_len]


# ----------------------------
# Sheet -> Matrix (fill merged)
# ----------------------------
def read_sheet_matrix_fill_merged(ws) -> List[List[str]]:
    """
    Read sheet into matrix of strings.
    For merged cells, fill the merged rectangle with the top-left value.
    Images are ignored automatically (we never touch them).
    """
    max_row = ws.max_row or 1
    max_col = ws.max_column or 1

    mat: List[List[Any]] = []
    for r in range(1, max_row + 1):
        row_vals = []
        for c in range(1, max_col + 1):
            row_vals.append(ws.cell(row=r, column=c).value)
        mat.append(row_vals)

    # Fill merged cells with top-left value
    for merged in ws.merged_cells.ranges:
        min_row, min_col, max_row2, max_col2 = merged.min_row, merged.min_col, merged.max_row, merged.max_col
        top_left = mat[min_row - 1][min_col - 1]
        for rr in range(min_row - 1, max_row2):
            for cc in range(min_col - 1, max_col2):
                mat[rr][cc] = top_left

    # Trim to used bounds
    r0, r1, c0, c1 = sheet_used_bounds(mat)
    if r1 < r0 or c1 < c0:
        return []

    trimmed: List[List[str]] = []
    for r in range(r0, r1 + 1):
        trimmed.append([normalize_cell(mat[r][c]) for c in range(c0, c1 + 1)])
    return trimmed


# ----------------------------
# Matrix -> Blocks (split by blank rows)
# ----------------------------
def split_blocks_by_blank_rows(matrix: List[List[str]], blank_rows: int = 1) -> List[List[List[str]]]:
    """
    Split matrix into blocks by consecutive blank rows (row where all cells empty strings).
    blank_rows: number of consecutive blank rows required to split.
    """
    if not matrix:
        return []

    def row_is_blank(row: List[str]) -> bool:
        return all((cell.strip() == "") for cell in row)

    blocks: List[List[List[str]]] = []
    current: List[List[str]] = []
    blank_run = 0

    for row in matrix:
        if row_is_blank(row):
            blank_run += 1
            if current and blank_run >= blank_rows:
                blocks.append(current)
                current = []
        else:
            blank_run = 0
            current.append(row)

    if current:
        blocks.append(current)
    return blocks


def trim_empty_cols(block: List[List[str]]) -> List[List[str]]:
    """Trim leading/trailing columns that are entirely empty within this block."""
    if not block:
        return block
    ncols = max(len(r) for r in block)
    padded = [r + [""] * (ncols - len(r)) for r in block]

    non_empty_cols = []
    for c in range(ncols):
        col_has = any(padded[r][c].strip() != "" for r in range(len(padded)))
        if col_has:
            non_empty_cols.append(c)

    if not non_empty_cols:
        return []

    c0, c1 = min(non_empty_cols), max(non_empty_cols)
    return [row[c0 : c1 + 1] for row in padded]


def block_to_markdown_table(block: List[List[str]]) -> str:
    if not block:
        return ""
    block = trim_empty_cols(block)
    if not block:
        return ""

    ncols = max(len(r) for r in block)
    padded = [r + [""] * (ncols - len(r)) for r in block]

    # ヘッダなし要件：Col1.. を付与して全行データ扱い
    columns = [f"Col{i+1}" for i in range(ncols)]
    df = pd.DataFrame(padded, columns=columns)
    return df.to_markdown(index=False, tablefmt="github")


# ----------------------------
# DOCX processing functions
# ----------------------------
def extract_table_from_docx_table(table: Table) -> List[List[str]]:
    """Extract table data from docx Table object."""
    rows = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            # セル内のテキストを取得（改行は空白に置換）
            cell_text = cell.text.strip().replace('\n', ' ').replace('\r', ' ')
            cells.append(cell_text)
        rows.append(cells)
    return rows


def convert_docx_to_md(docx_path: Path, out_dir: Path, blank_rows: int = 1) -> List[Path]:
    """Convert a DOCX file to Markdown format."""
    doc = Document(docx_path)
    out_dir.mkdir(parents=True, exist_ok=True)
    
    base = docx_path.stem
    out_path = out_dir / f"{base}.md"
    
    parts: List[str] = []
    parts.append(f"# {docx_path.name}")
    parts.append("")
    
    table_count = 0
    
    for element in doc.element.body:
        # 段落の処理
        if element.tag.endswith('p'):
            para = None
            for p in doc.paragraphs:
                if p._element == element:
                    para = p
                    break
            
            if para:
                text = para.text.strip()
                if not text:
                    continue
                
                # 見出しスタイルの判定
                style_name = para.style.name if para.style else ""
                if "Heading" in style_name:
                    # 見出しレベルを抽出（Heading 1 -> #, Heading 2 -> ##, など）
                    level_match = re.search(r'Heading (\d+)', style_name)
                    if level_match:
                        level = int(level_match.group(1))
                        parts.append(f"{'#' * (level + 1)} {text}")
                    else:
                        parts.append(f"## {text}")
                else:
                    # 通常の段落
                    parts.append(text)
                parts.append("")
        
        # 表の処理
        elif element.tag.endswith('tbl'):
            table = None
            for t in doc.tables:
                if t._element == element:
                    table = t
                    break
            
            if table:
                table_count += 1
                table_data = extract_table_from_docx_table(table)
                
                if table_data:
                    # 空の行や列を除去
                    blocks = split_blocks_by_blank_rows(table_data, blank_rows=max(1, blank_rows))
                    
                    if len(blocks) == 1:
                        md = block_to_markdown_table(blocks[0])
                        if md.strip():
                            parts.append(f"### Table {table_count}")
                            parts.append("")
                            parts.append(md)
                            parts.append("")
                    else:
                        sub_table_no = 0
                        for block in blocks:
                            md = block_to_markdown_table(block)
                            if not md.strip():
                                continue
                            sub_table_no += 1
                            parts.append(f"### Table {table_count}-{sub_table_no}")
                            parts.append("")
                            parts.append(md)
                            parts.append("")
    
    if not any(part.strip() for part in parts[2:]):  # ヘッダー以外に内容がない場合
        parts.append("_（このドキュメントは空、または内容が見つかりませんでした）_")
    
    out_path.write_text("\n".join(parts).rstrip() + "\n", encoding="utf-8")
    return [out_path]


# ----------------------------
# Workbook -> (per-sheet md files)
# ----------------------------
def convert_excel_to_md_per_sheet(xlsx_path: Path, out_dir: Path, blank_rows: int = 1) -> List[Path]:
    wb = load_workbook(xlsx_path, data_only=True)  # 表示値（保存済みの計算結果）を読む
    out_dir.mkdir(parents=True, exist_ok=True)

    created: List[Path] = []
    base = xlsx_path.stem

    for idx, sheet_name in enumerate(wb.sheetnames, start=1):
        ws = wb[sheet_name]
        matrix = read_sheet_matrix_fill_merged(ws)

        safe_sheet = safe_filename(sheet_name)
        out_path = out_dir / f"{base}__{idx:02d}_{safe_sheet}.md"

        parts: List[str] = []
        parts.append(f"# {xlsx_path.name}")
        parts.append(f"## {sheet_name}")
        parts.append("")

        if not matrix:
            parts.append("_（このシートは空、または値が見つかりませんでした）_")
            out_path.write_text("\n".join(parts) + "\n", encoding="utf-8")
            created.append(out_path)
            continue

        blocks = split_blocks_by_blank_rows(matrix, blank_rows=max(1, blank_rows))

        if len(blocks) == 1:
            md = block_to_markdown_table(blocks[0])
            parts.append(md if md.strip() else "_（表が生成できませんでした）_")
        else:
            table_no = 0
            for block in blocks:
                md = block_to_markdown_table(block)
                if not md.strip():
                    continue
                table_no += 1
                parts.append(f"### Table {table_no}")
                parts.append("")
                parts.append(md)
                parts.append("")

            if table_no == 0:
                parts.append("_（表が生成できませんでした）_")

        out_path.write_text("\n".join(parts).rstrip() + "\n", encoding="utf-8")
        created.append(out_path)

    return created


# ----------------------------
# Directory traversal
# ----------------------------
def convert_tree(input_root: Path, output_root: Path, blank_rows: int = 1) -> Tuple[int, int]:
    """
    Convert all .xlsx and .docx files under input_root into output_root, preserving subdirectory structure.
    Returns: (ok_count, ng_count) for documents
    """
    ok = 0
    ng = 0

    # Process Excel files (.xlsx)
    xlsx_files = sorted(input_root.rglob("*.xlsx"))
    for xlsx in xlsx_files:
        # Excel作業中の一時ファイルをスキップ
        if xlsx.name.startswith("~$"):
            continue

        rel_dir = xlsx.parent.relative_to(input_root)
        out_dir = output_root / rel_dir

        try:
            convert_excel_to_md_per_sheet(xlsx, out_dir, blank_rows=blank_rows)
            ok += 1
            print(f"[OK] {xlsx} -> {out_dir}")
        except Exception as e:
            ng += 1
            print(f"[NG] {xlsx} : {e}", file=sys.stderr)

    # Process Word files (.docx)
    docx_files = sorted(input_root.rglob("*.docx"))
    for docx in docx_files:
        # Word作業中の一時ファイルをスキップ
        if docx.name.startswith("~$"):
            continue

        rel_dir = docx.parent.relative_to(input_root)
        out_dir = output_root / rel_dir

        try:
            convert_docx_to_md(docx, out_dir, blank_rows=blank_rows)
            ok += 1
            print(f"[OK] {docx} -> {out_dir}")
        except Exception as e:
            ng += 1
            print(f"[NG] {docx} : {e}", file=sys.stderr)

    return ok, ng


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--in", dest="in_dir", default="設計書", help='Input root directory (default: "設計書")')
    ap.add_argument("--out", dest="out_dir", default="converted_md_design", help='Output root directory (default: "converted_md_design")')
    ap.add_argument("--blank-rows", type=int, default=1, help="Consecutive blank rows to split tables (default: 1)")
    args = ap.parse_args()

    input_root = Path(args.in_dir).resolve()
    output_root = Path(args.out_dir).resolve()

    if not input_root.exists():
        print(f'Input directory not found: {input_root}  （repo直下に「設計書」フォルダを作ってxlsxやdocxファイルを置いてください）', file=sys.stderr)
        return 2

    ok, ng = convert_tree(input_root, output_root, blank_rows=max(1, args.blank_rows))
    print(f"\nDone. documents OK={ok}, NG={ng}")
    return 1 if ng > 0 else 0


if __name__ == "__main__":
    raise SystemExit(main())
